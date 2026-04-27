"""
ResumeParser: extracts plain text from PDF, DOCX, or TXT resumes.

Design priorities:
  - graceful degradation: if a parser fails, fall back to the next
  - cleaning: remove control chars, normalize whitespace, fix word splits
    that PDF parsers commonly introduce ("Pyt hon" -> "Python")
  - layout probe: surface ATS-unfriendly structures (tables, multi-column
    layouts) for the scorer
"""

from __future__ import annotations

import logging
import os
import re
import tempfile
from typing import Dict, Optional

logger = logging.getLogger(__name__)


class ResumeParser:
    """Extract raw text from a resume file."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

    def __init__(self) -> None:
        # Lazy-import heavy libs.
        self._pdfplumber = None
        self._pypdf2 = None
        self._docx = None

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def parse(self, file_path: str) -> str:
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported extension '{ext}'. "
                f"Supported: {sorted(self.SUPPORTED_EXTENSIONS)}"
            )

        if ext == ".pdf":
            raw = self._parse_pdf(file_path)
        elif ext == ".docx":
            raw = self._parse_docx(file_path)
        else:
            raw = self._parse_txt(file_path)

        return self._clean(raw)

    def parse_from_bytes(self, data: bytes, extension: str) -> str:
        ext = extension.lower()
        if not ext.startswith("."):
            ext = "." + ext
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported extension '{ext}'.")

        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        try:
            return self.parse(tmp_path)
        finally:
            try:
                os.remove(tmp_path)
            except OSError:
                pass

    # ------------------------------------------------------------------
    # Format-specific parsers
    # ------------------------------------------------------------------

    def _parse_pdf(self, file_path: str) -> str:
        # pdfplumber is more accurate for layout; fall back to PyPDF2.
        text = self._try_pdfplumber(file_path)
        if text and text.strip():
            return text
        logger.info("pdfplumber returned empty text; trying PyPDF2.")
        return self._try_pypdf2(file_path)

    def _try_pdfplumber(self, file_path: str) -> str:
        try:
            if self._pdfplumber is None:
                import pdfplumber
                self._pdfplumber = pdfplumber
            chunks = []
            with self._pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text(
                        x_tolerance=2, y_tolerance=2
                    ) or ""
                    chunks.append(page_text)
            return "\n".join(chunks)
        except Exception as exc:  # noqa: BLE001
            logger.warning("pdfplumber failed on %s: %s", file_path, exc)
            return ""

    def _try_pypdf2(self, file_path: str) -> str:
        try:
            if self._pypdf2 is None:
                import PyPDF2
                self._pypdf2 = PyPDF2
            chunks = []
            with open(file_path, "rb") as fh:
                reader = self._pypdf2.PdfReader(fh)
                for page in reader.pages:
                    try:
                        chunks.append(page.extract_text() or "")
                    except Exception as page_exc:  # noqa: BLE001
                        logger.warning("PyPDF2 page error: %s", page_exc)
            return "\n".join(chunks)
        except Exception as exc:  # noqa: BLE001
            logger.error("PyPDF2 failed on %s: %s", file_path, exc)
            return ""

    def _parse_docx(self, file_path: str) -> str:
        try:
            if self._docx is None:
                import docx
                self._docx = docx
            document = self._docx.Document(file_path)
            paragraphs = [p.text for p in document.paragraphs if p.text]
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            paragraphs.append(cell_text)
            return "\n".join(paragraphs)
        except Exception as exc:  # noqa: BLE001
            logger.error("python-docx failed on %s: %s", file_path, exc)
            return ""

    @staticmethod
    def _parse_txt(file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as fh:
                return fh.read()
        except Exception as exc:  # noqa: BLE001
            logger.error("Failed to read txt %s: %s", file_path, exc)
            return ""

    # ------------------------------------------------------------------
    # Cleaning — fix common PDF extraction artifacts
    # ------------------------------------------------------------------

    @staticmethod
    def _clean(text: str) -> str:
        if not text:
            return ""
        # Normalize line endings
        text = re.sub(r"\r\n?", "\n", text)
        # Strip control chars except newline/tab
        text = re.sub(r"[\x00-\x08\x0b-\x1f\x7f]", "", text)
        # Replace common ligatures and smart quotes
        replacements = {
            "\ufb00": "ff", "\ufb01": "fi", "\ufb02": "fl",
            "\ufb03": "ffi", "\ufb04": "ffl",
            "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"',
            "\u2013": "-", "\u2014": "-", "\u2022": "*",
        }
        for src, tgt in replacements.items():
            text = text.replace(src, tgt)
        # Glue back single-letter splits PDFs sometimes produce: "P y t h o n"
        # Heuristic: if we see 4+ single letters separated by single spaces,
        # squash them.
        text = re.sub(
            r"\b(?:[A-Za-z] ){3,}[A-Za-z]\b",
            lambda m: m.group(0).replace(" ", ""),
            text,
        )
        # Collapse excessive whitespace
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    # ------------------------------------------------------------------
    # Layout-quality probe (used by ATSScorer)
    # ------------------------------------------------------------------

    def detect_layout_issues(self, file_path: Optional[str]) -> Dict:
        flags = {"has_tables": False, "page_count": 0,
                 "has_multiple_columns": False}
        if not file_path or not os.path.isfile(file_path):
            return flags
        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext == ".docx":
                if self._docx is None:
                    import docx
                    self._docx = docx
                document = self._docx.Document(file_path)
                flags["has_tables"] = len(document.tables) > 0
            elif ext == ".pdf":
                if self._pdfplumber is None:
                    import pdfplumber
                    self._pdfplumber = pdfplumber
                with self._pdfplumber.open(file_path) as pdf:
                    flags["page_count"] = len(pdf.pages)
                    for page in pdf.pages:
                        if page.find_tables():
                            flags["has_tables"] = True
                        # Heuristic for multi-column layout: look at the
                        # x-coords of word starts; a strong bimodal pattern
                        # suggests two columns.
                        words = page.extract_words() or []
                        if len(words) > 30:
                            xs = sorted(w["x0"] for w in words)
                            page_w = page.width or 612
                            left_count = sum(1 for x in xs if x < page_w / 2)
                            right_count = len(xs) - left_count
                            ratio = (
                                min(left_count, right_count) /
                                max(left_count, right_count, 1)
                            )
                            # If both halves contain substantial text,
                            # likely multi-column.
                            if ratio > 0.4 and right_count > 15:
                                flags["has_multiple_columns"] = True
        except Exception as exc:  # noqa: BLE001
            logger.warning("Layout probe failed: %s", exc)
        return flags
